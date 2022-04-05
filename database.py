import sqlite3

from docx import Document
from docx2txt import docx2txt


def delete():
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()

    # ----------------------ОК Базовые--------------------
    cursor.execute("""DROP TABLE IF EXISTS ok_base""")

    # --------------------ОК Углубленные------------------
    cursor.execute("""DROP TABLE IF EXISTS ok_immersed""")

    # ----------------------ПК Базовые--------------------
    cursor.execute("""DROP TABLE IF EXISTS pk_base""")

    # --------------------ПК Углубленные------------------
    cursor.execute("""DROP TABLE IF EXISTS pk_immersed""")

    # -----------------Компетенции Углубленные------------
    cursor.execute("""DROP TABLE IF EXISTS competencies_base""")

    # -----------------Компетенции Базовые----------------
    cursor.execute("""DROP TABLE IF EXISTS competencies_immersed""")

    connect.commit()


def xao(string):
    slash = ' \ '
    return string.rstrip().replace('\n', '').replace(slash.replace(' ', '') + 'xc2', '').replace(
        slash.replace(' ', '') + 'xa0', '')


def create():
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    # ----------------------ОК Базовые--------------------
    cursor.execute("""CREATE TABLE IF NOT EXISTS ok_base(
        ok TEXT,
        ok_text TEXT
    )""")
    # --------------------ОК Углубленные------------------
    cursor.execute("""CREATE TABLE IF NOT EXISTS ok_immersed(
        ok TEXT,
        ok_text TEXT
    )""")
    # ----------------------ПК Базовые--------------------
    cursor.execute("""CREATE TABLE IF NOT EXISTS pk_base(
        pk TEXT,
        pk_text TEXT
    )""")
    # --------------------ПК Углубленные------------------
    cursor.execute("""CREATE TABLE IF NOT EXISTS pk_immersed(
        pk TEXT,
        pk_text TEXT
    )""")
    # -----------------Компетенции Углубленные------------
    cursor.execute("""CREATE TABLE IF NOT EXISTS competencies_base(
        competencies TEXT,
        ok TEXT,
        pk TEXT
    )""")
    # -----------------Компетенции Базовые----------------
    cursor.execute("""CREATE TABLE IF NOT EXISTS competencies_immersed(
        competencies TEXT,
        ok TEXT,
        pk TEXT
    )""")

    connect.commit()


def update_ok(document_name):
    list_ok = []
    need_text = ''
    can = True
    for line in docx2txt.process(document_name):
        if line == '\n':
            if need_text[0:2] == 'ОК':  # Rus
                head = 'ОК '  # Rus
                for letter in range(3, len(need_text), 2):
                    if need_text[letter + 1] == '.':
                        head += need_text[letter] + need_text[letter + 1]
                    else:
                        break
                if head == 'ОК ':  # Rus
                    need_text = ''
                    continue
                else:
                    info = need_text[len(head):]
                    if list_ok.count(head) == 0 and can:
                        try:
                            list_ok.append(head)
                            list_ok.append(info)
                        except:
                            print('Не удалось записать данные в ok_base')
                            break
                    else:
                        if can:
                            save_db('ok_base', list_ok, 2)
                            list_ok = []
                        can = False
                        try:
                            list_ok.append(head)
                            list_ok.append(info)
                        except:
                            print('Не удалось записать данные в ok_immersed')
                            break
            if need_text[:7] == 'Таблица' and not can:
                save_db('ok_immersed', list_ok, 2)
                break
            need_text = ''
        else:
            need_text += line


def update_pk(document_name):
    list_pk = []
    need_text = ''
    can = True
    for line in docx2txt.process(document_name):
        if line == '\n':
            if need_text[0:2] == 'ПК':  # Rus
                head = 'ПК '  # Rus
                for letter in range(3, len(need_text), 2):
                    if need_text[letter + 1] == '.':
                        head += need_text[letter] + need_text[letter + 1]
                    else:
                        break
                if head == 'ПК ':  # Rus
                    need_text = ''
                    continue
                else:
                    info = need_text[len(head):]
                    if list_pk.count(head) == 0 and can:
                        try:
                            list_pk.append(head)
                            list_pk.append(info)
                        except:
                            print('Не удалось записать данные в pk_base')
                            break
                    else:
                        if can:
                            save_db('pk_base', list_pk, 2)
                            list_pk = []
                        can = False
                        try:
                            list_pk.append(head)
                            list_pk.append(info)
                        except:
                            print('Не удалось записать данные в pk_immersed')
                            break
            if need_text[:7] == 'Таблица' and not can:
                save_db('pk_immersed', list_pk, 2)
                break
            need_text = ''
        else:
            need_text += line


def try_int(mb_num):
    try:
        int(mb_num)
        return True
    except:
        return False


def find_twins(cell):
    competencies = ''
    twin = ''
    twins = []
    for letter in range(0, len(cell)):
        if try_int(cell[letter]):
            break
        competencies += cell[letter]
    for letter in range(0, len(cell)):
        if cell[letter:len(competencies)] == competencies:
            twin += cell[letter:len(competencies)]
            for name in range(len(cell[letter:len(competencies)]), len(cell)):
                if cell[name:name + len(competencies)] == competencies:
                    twins.append(xao(str(twin)))
                    twin = ''
                twin += cell[name]
                if name == len(cell) - 1:
                    twin.rstrip()
                    twins.append(xao(str(twin)))
            break
    return twins


def try_ten(letter, need_text, need_num, were):
    if len(need_text) > letter:
        if were < 0:
            n = -1
        else:
            n = len(need_text)
        for i in range(letter, n, were):
            if try_int(need_text[i]) or need_text[i] == ' ':
                need_num += need_text[i]
                if were < 0:
                    if i == 0:
                        return need_num
                else:
                    if i == len(need_text) - 1:
                        return need_num
            else:
                return need_num

    else:
        return need_num


def form_competencies(find, twin, document_name, immersed):
    findling = find[twin]
    need_text = ''
    can = False
    competencies = ''
    need_to_needling = 2
    for line in docx2txt.process(document_name):
        if line == '\n':
            if need_text == 'Таблица 4':
                immersed = False
            if not immersed:
                if len(find) > 1:
                    if need_text[:len(find[0])] == find[0]:
                        for i in range(0, len(need_text)):
                            if need_text[i:len(findling) + i] == findling:
                                can = True
                if need_text == findling:
                    can = True
                if can:
                    if need_to_needling == 0:
                        competencies = []
                        need_num = ''
                        tag = ''
                        carelessness = True
                        for letter in range(0, len(need_text)):
                            if letter == len(need_text) - 1:
                                competencies.append(tag + need_num)
                            elif need_text[letter] + need_text[letter + 1] == 'ОК':  # Rus
                                if need_num != '' and tag != 'ОК':
                                    competencies.append(tag + need_num)
                                tag = 'ОК' + need_text[letter + 2]
                                need_num = ''
                            elif need_text[letter] + need_text[letter + 1] == 'ПК':  # Rus
                                if need_num != '' and tag != 'ПК':
                                    competencies.append(tag + need_num)
                                tag = 'ПК' + need_text[letter + 2]
                                need_num = ''
                            elif need_text[letter] == '.':
                                if carelessness:
                                    need_num = try_ten(letter - 2, need_text, need_num, - 1)
                                    need_num += need_text[letter - 1:letter + 2]
                                    need_num = try_ten(letter + 2, need_text, need_num, + 1)
                                else:
                                    need_num += ', ' + need_text[letter - 1:letter + 2]
                                carelessness = False
                            elif need_text[letter] == '-':
                                if carelessness:
                                    need_num = try_ten(letter - 3, need_text, need_num, - 1)
                                    need_num += need_text[letter - 2:letter + 3]
                                    need_num = try_ten(letter + 3, need_text, need_num, + 1)
                                else:
                                    need_num += ' ' + need_text[letter] + ' '
                                carelessness = True
                            elif need_text[letter] == ',':
                                if carelessness:
                                    need_num = try_ten(letter - 2, need_text, need_num, - 1)
                                    need_num += need_text[letter - 1:letter + 1] + ' '
                                    if len(need_text) == letter + 3:
                                        try:
                                            need_num += need_text[letter + 2]
                                            need_num = try_ten(letter + 3, need_text, need_num, + 1)
                                        except:
                                            continue
                                else:
                                    need_num += need_text[letter] + ' '
                                carelessness = True
                        break
                    need_to_needling -= 1
            need_text = ''
        else:
            need_text += line
    return competencies, findling


def update_competencies_base(document_name):
    document = Document(document_name)
    for line in document.tables[2].rows:
        if line.cells[4].text.strip() != '' and line.cells[4].text != document.tables[4].rows[0].cells[4].text.strip():
            find = []
            find.extend(find_twins(line.cells[4].text))
            for twin in range(0, len(find)):
                competencies, findling = form_competencies(find, twin, document_name, False)
                if competencies != '':
                    competencies_list = [findling]
                    for element in competencies:
                        competencies_list.append(element[:2] + ' ' + element[2:])
                    if len(competencies_list) < 3:
                        for i in range(0, 3 - len(competencies_list)):
                            competencies_list.append('')
                    save_db('competencies_base', competencies_list, 3)


def update_competencies_immersed(document_name):
    document = Document(document_name)
    for line in document.tables[4].rows:
        if line.cells[4].text.strip() != '' and line.cells[4].text != document.tables[4].rows[0].cells[4].text.strip():
            find = []
            find.extend(find_twins(line.cells[4].text))
            for twin in range(0, len(find)):
                competencies, findling = form_competencies(find, twin, document_name, True)
                if competencies != '':
                    competencies_list = [findling]
                    for element in competencies:
                        competencies_list.append(element[:2] + ' ' + element[2:])
                    if len(competencies_list) < 3:
                        for i in range(0, 3 - len(competencies_list)):
                            competencies_list.append('')
                    save_db('competencies_immersed', competencies_list, 3)


def give(need, findling, immersed):
    if immersed:
        if need == 'OK':
            return give_ok_immersed(findling)
        elif need == 'PK':
            return give_pk_immersed(findling)
    else:
        if need == 'OK':
            return give_ok_base(findling)
        elif need == 'PK':
            return give_pk_base(findling)


def calc(result, findling, tag):
    competencies = []
    n = 0
    if tag == 'ОК':
        n = 1
    elif tag == 'ПК':
        n = 2
    if n != 0:
        for directory in range(0, len(result)):
            if findling == result[directory][0]:
                dat = False
                ok = ''
                for nums in range(3, len(result[directory][n])):
                    if result[directory][n][nums] == '.':
                        ok = ''
                        dat = True
                        ok = try_ten(nums - 1, result[directory][n], ok, - 1).replace(' ', '')
                        ok += '.'
                        ok = try_ten(nums + 1, result[directory][n], ok, + 1).replace(' ', '')
                    elif result[directory][n][nums] == '-':
                        start = ''
                        finish = ''
                        if dat:
                            for i in range(nums, len(result[directory][n])):
                                if result[directory][n][i] == '.':
                                    finish = try_ten(i + 1, result[directory][n], finish, + 1)
                                    break
                            start = try_ten(nums - 1, result[directory][n], start, - 1)
                        else:
                            finish = try_ten(nums + 1, result[directory][n], finish, + 1)
                            start = try_ten(nums - 1, result[directory][n], start, - 1)
                        for i in range(int(start.replace(' ', '')), int(finish.replace(' ', '')) + 1):
                            competencies.append(tag + ' ' + (ok[:2]).rstrip() + str(i) + '.')
                        ok = ''
                    elif result[directory][n][nums] == ',':
                        if ok != '':
                            if tag + ' ' + ok + '.' != competencies[len(competencies) - 1]:
                                competencies.append(tag + ' ' + ok)
                                ok = ''
                break
    return competencies


def give_ok_immersed(findling):
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM competencies_immersed")
    result = cursor.fetchall()
    competencies = calc(result, findling, 'ОК')

    connect.commit()
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM ok_immersed")
    result = cursor.fetchall()
    out = []
    for i in range(0, len(competencies)):
        for j in range(0, len(result)):
            if result[j][0] == competencies[i]:
                out.append((result[j][0], result[j][1]))
    connect.commit()
    return out


def give_pk_immersed(findling):
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM competencies_immersed")
    result = cursor.fetchall()
    competencies = calc(result, findling, 'ПК')
    connect.commit()
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM pk_immersed")
    result = cursor.fetchall()
    out = []
    for i in range(0, len(competencies)):
        for j in range(0, len(result)):
            if result[j][0] == competencies[i]:
                out.append((result[j][0], result[j][1]))
    connect.commit()
    return out


def give_pk_base(findling):
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM competencies_base")
    result = cursor.fetchall()
    competencies = calc(result, findling, 'ПК')
    connect.commit()
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM pk_base")
    result = cursor.fetchall()
    out = []
    for i in range(0, len(competencies)):
        for j in range(0, len(result)):
            if result[j][0] == competencies[i]:
                out.append((result[j][0], result[j][1]))
    connect.commit()
    return out


def give_ok_base(findling):
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM competencies_base")
    result = cursor.fetchall()
    competencies = calc(result, findling, 'ОК')
    connect.commit()
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    cursor.execute("SELECT * FROM ok_base")
    result = cursor.fetchall()
    out = []
    for i in range(0, len(competencies)):
        for j in range(0, len(result)):
            if result[j][0] == competencies[i]:
                out.append((result[j][0], result[j][1]))
    connect.commit()
    return out


def save_db(name_table, list, line):
    connect = sqlite3.connect('competencies.db')
    cursor = connect.cursor()
    needling = ''
    for i in range(0, line):
        if i == line - 1:
            needling += '?'
            break
        needling += '?,'
    for i in range(0, len(list), line):
        save = []
        for j in range(0, line):
            save.append(list[j + i])
        cursor.execute("INSERT INTO " + name_table + " VALUES(" + needling + ");", save)
    connect.commit()


if __name__ == '__main__':
    print(give('OK', 'МДК.03.03. Документирование и сертификация', False))
    print(give('PK', 'МДК.03.03. Документирование и сертификация', False))
